from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.conf import settings
import re

FIXED_PASSWORD = 'contraseña_segura'  # Reemplaza con tu contraseña fija


def login_view(request):
    if request.method == 'POST':
        password = request.POST.get('password')
        if password == FIXED_PASSWORD:
            # Autenticar al usuario
            user = authenticate(request, username='admin', password=FIXED_PASSWORD)
            if user is None:
                # Crear un usuario si no existe
                from django.contrib.auth.models import User
                user = User.objects.create_user('admin', password=FIXED_PASSWORD)
            login(request, user)
            return redirect('dashboard')
        else:
            messages.error(request, 'Contraseña incorrecta.')
    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

def dashboard_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el texto del PDF y las secciones almacenadas en la sesión
    pdf_text = request.session.get('pdf_text', '')
    asignatura = request.session.get('Asignatura', 'No especificada')
    titulacion = request.session.get('Titulación', 'No especificada')
    facultad = request.session.get('Facultad/Escuela', 'No especificada')

    return render(request, 'dashboard.html', {
        'pdf_text': pdf_text,
        'asignatura': asignatura,
        'titulacion': titulacion,
        'facultad': facultad
    })




# Incorporamos la subida de archivos PDF
import os
from django.conf import settings
from django.core.files.storage import FileSystemStorage
import PyPDF2

def upload_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    if request.method == 'POST' and request.FILES.get('pdf_file'):
        pdf_file = request.FILES['pdf_file']
        try:
            fs = FileSystemStorage(location=settings.MEDIA_ROOT)
            filename = fs.save(pdf_file.name, pdf_file)

            # Procesar el PDF para extraer el texto
            pdf_path = os.path.join(settings.MEDIA_ROOT, filename)
            text_content = extract_text_from_pdf(pdf_path)

            # Guardar el texto en la sesión para usarlo en otras vistas
            request.session['pdf_text'] = text_content

            # Almacenar las secciones extraídas en la sesión
            sections = process_and_parse_pdf_text(text_content)

            # Guardar las secciones principales en la sesión
            request.session['Asignatura'] = sections.get('Asignatura', 'No especificada')
            request.session['Titulación'] = sections.get('Titulación', 'No especificada')
            request.session['Facultad/Escuela'] = sections.get('Facultad/Escuela', 'No especificada')
            request.session['CONTENIDOS'] = sections.get('CONTENIDOS', '')
            request.session['temas'] = sections.get('temas', [])
            
            return render(request, 'pdf_text_view.html', {'text_content': text_content})
        except Exception as e:
            print(f"Error al procesar el archivo: {e}")
            return render(request, 'upload.html', {'error_message': 'Error al procesar el archivo PDF. Verifica el formato.'})

    return render(request, 'upload.html')





def extract_text_from_pdf(pdf_path):
    text_content = ''
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text_content += page.extract_text()
    
    # Eliminar los saltos de línea del texto
    text_content = text_content.replace('\n', ' ')
    
    # Verificar si el texto del PDF se está extrayendo correctamente
    # print(f"Texto extraído del PDF: {text_content[:2000]}")  # Muestra los primeros 500 caracteres
    return text_content


def process_and_parse_pdf_text(pdf_text):
    # Limpiar el texto eliminando los datos personales y referencias de páginas
    cleaned_text = re.sub(r'Equipo Docente Correo Electrónico.*?DESCRIPCIÓN DE LA ASIGNATURA', 'DESCRIPCIÓN DE LA ASIGNATURA', pdf_text, flags=re.DOTALL)
    cleaned_text = re.sub(r'Página \d+', '', cleaned_text)

    # Prompt para estructurar el texto y extraer secciones
    prompt = (
        f"Estructura el siguiente texto en categorías claras como: "
        f"Titulación, Facultad/Escuela, Asignatura, Periodo docente, Tipo de enseñanza, "
        f"Idioma, DESCRIPCIÓN DE LA ASIGNATURA, OBJETIVO, CONOCIMIENTOS PREVIOS, "
        f"CONTENIDOS (incluyendo temas separados por '&&'), ACTIVIDADES FORMATIVAS, "
        f"DISTRIBUCIÓN DE LOS TIEMPOS DE TRABAJO, COMPETENCIAS, "
        f"RESULTADOS DE APRENDIZAJE, SISTEMA DE EVALUACIÓN DEL APRENDIZAJE, "
        f"BIBLIOGRAFÍA Y OTROS RECURSOS. "
        f"Por favor, usa '||' para delimitar las secciones y '&&' para los temas dentro de CONTENIDOS.\n\n"
        f"Texto:\n{cleaned_text}"
    )

    # Llamar a la API de la IA para estructurar el texto
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=2000,
            temperature=0.0,
        ),
    )

    structured_text = response.text

    # Proceso para dividir el texto estructurado en secciones usando '||' y '&&'
    section_titles = [
        'Titulación', 'Facultad/Escuela', 'Asignatura', 'Periodo docente', 'Tipo de enseñanza', 
        'Idioma', 'DESCRIPCIÓN DE LA ASIGNATURA', 'OBJETIVO', 'CONOCIMIENTOS PREVIOS', 
        'CONTENIDOS', 'ACTIVIDADES FORMATIVAS', 'DISTRIBUCIÓN DE LOS TIEMPOS DE TRABAJO', 
        'COMPETENCIAS', 'RESULTADOS DE APRENDIZAJE', 'SISTEMA DE EVALUACIÓN DEL APRENDIZAJE', 
        'BIBLIOGRAFÍA Y OTROS RECURSOS'
    ]

    # Eliminar el formato markdown (los asteriscos '**')
    structured_text = structured_text.replace('**', '')

    # Dividir el texto en partes utilizando "||" como delimitador y limpiar espacios
    parts = [part.strip() for part in structured_text.split('||') if part.strip()]

    # Diccionario para almacenar secciones y lista para temas
    sections = {}
    current_section = None
    themes = []  # Lista para almacenar los temas de la sección CONTENIDOS
    
    for part in parts:
        if part in section_titles:
            # Si encontramos un título de sección, lo guardamos como sección actual
            current_section = part
            sections[current_section] = ''
        elif current_section == 'CONTENIDOS':
            # Si estamos en la sección CONTENIDOS, verificamos si hay temas delimitados por '&&'
            if '&&' in part:
                # Dividimos los temas por el delimitador '&&' y los agregamos a la lista de temas
                themes = [theme.strip() for theme in part.split('&&') if theme.strip()]
            sections[current_section] += part + " "
        elif current_section:
            # Agregar el contenido a la sección actual
            sections[current_section] += part + " "

    # **Expresiones regulares para capturar las secciones faltantes** (como Titulación, Facultad, Asignatura)
    if 'Asignatura' not in sections or not sections['Asignatura'].strip():
        # Capturamos hasta "Tipo:", ya que es el siguiente campo
        match = re.search(r'Asignatura:\s*(.*?)(?=\s*Tipo:)', pdf_text, re.DOTALL)
        if match:
            sections['Asignatura'] = match.group(1).strip()

    if 'Titulación' not in sections or not sections['Titulación'].strip():
        # Capturamos hasta "Rama de Conocimiento:", ya que es el siguiente campo
        match = re.search(r'Titulación:\s*(.*?)(?=\s*Rama de Conocimiento:)', pdf_text, re.DOTALL)
        if match:
            sections['Titulación'] = match.group(1).strip()

    if 'Facultad/Escuela' not in sections or not sections['Facultad/Escuela'].strip():
        # Capturamos hasta "Asignatura:", ya que es el siguiente campo
        match = re.search(r'Facultad(?:/Escuela)?:\s*(.*?)(?=\s*Asignatura:)', pdf_text, re.DOTALL)
        if match:
            sections['Facultad/Escuela'] = match.group(1).strip()

    sections['temas'] = themes
    print(f"temas: {themes}")   

    return sections



def extract_description(pdf_text):
    # Eliminar referencias de páginas como "Página 1", "Página 2"
    pdf_text = re.sub(r'Página \d+', '', pdf_text)

    # Buscar "DESCRIPCIÓN DE LA ASIGNATURA" hasta "OBJETIVO"
    match = re.search(r"DESCRIPCIÓN DE LA ASIGNATURA(.*?)(OBJETIVO)", pdf_text, re.DOTALL)
    if match:
        return match.group(1).strip()  # Devolver el texto entre ambos títulos
    return None


                          




# Añadimos la integración con GEMINI 
import google.generativeai as genai
# Configurar la clave de API
api_key = os.getenv('API_KEY_GEMINI_GUIAS_DOCENTES')
if api_key is None:
    print("Error: La variable de entorno API_KEY_GEMINI_GUIAS_DOCENTES no está definida")
    # Aquí puedes agregar código para manejar el error
else:
    genai.configure(api_key=api_key) 

### AQUÍ EMPIEZAN LAS APLICACIONES

## REFORMULAR LA DESCRIPCIÓN

def reformulate_description(request):
    if not request.user.is_authenticated:
        return redirect('login')

    pdf_text = request.session.get('pdf_text', '')
    programa_formativo = load_programa_formativo()

    # Extraer el texto entre "DESCRIPCIÓN DE LA ASIGNATURA" y "OBJETIVO"
    description_text = extract_description(pdf_text)

    if not description_text:
        return render(request, 'reformulate_description.html', {'error_message': 'No se pudo encontrar la descripción de la asignatura.'})

    # Obtener el texto completo estructurado desde la sesión
    pdf_text = request.session.get('pdf_text', '')

    if not pdf_text:
        return render(request, 'expand_content.html', {'error_message': 'No se ha encontrado el texto del PDF.'})

    # Llamar a la función unificada para procesar y estructurar el contenido
    sections = process_and_parse_pdf_text(pdf_text)

    # Obtener las actividades formativas
    sistema_evaluacion = sections.get('SISTEMA DE EVALUACIÓN DEL APRENDIZAJE', '')
    resultados_aprendizaje = sections.get('RESULTADOS DE APRENDIZAJE', '')
    asignatura = sections.get('Asignatura', '')
    competencias = sections.get('COMPETENCIAS', '')
    contenidos = sections.get('CONTENIDOS', '')
    titulacion = sections.get('Titulación', '')
    facultad = sections.get('Facultad/Escuela', '')
    periodo_docente = sections.get('Periodo docente', '')
    programa_formativo = load_programa_formativo()

    prompt = f"Aquí tienes la DESCRIPCIÓN de la asignatura {asignatura} de una titulación universitaria ({titulacion} que pertenece a la facultad de {facultad}). Esto te lo doy como contexto, no para que hagas referencia expresa. Reformula la siguiente descripción de la asignatura incorporando los principios del modelo pedagógico:\n\n {programa_formativo}.\n\n La descripición no debe ser comercial; no se trata de vender, si no mostrar lo que se va a impartir (sin ser excesivamente formal). Los títulos no los escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas). \n\nDescripción actual:\n{description_text}"
    

    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=1500,
            temperature=0.4,
        ),
    )

    reformulated_text = response.text
    # print(f"Descripción reformulada: {reformulated_text}")

    # Renderizar el resultado con ambas descripciones
    return render(request, 'reformulate_description.html', {
        'original_description': description_text,
        'reformulated_description': reformulated_text
    })


## EXPANDIR EL CONTENIDO

def expand_content(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el texto completo estructurado desde la sesión
    pdf_text = request.session.get('pdf_text', '')

    if not pdf_text:
        return render(request, 'expand_content.html', {'error_message': 'No se ha encontrado el texto del PDF.'})

    # Llamar a la función unificada para procesar y estructurar el contenido
    sections = process_and_parse_pdf_text(pdf_text)

    # Obtener las actividades formativas
    sistema_evaluacion = sections.get('SISTEMA DE EVALUACIÓN DEL APRENDIZAJE', '')
    resultados_aprendizaje = sections.get('RESULTADOS DE APRENDIZAJE', '')
    asignatura = sections.get('Asignatura', '')
    competencias = sections.get('COMPETENCIAS', '')
    contenidos = sections.get('CONTENIDOS', '')
    titulacion = sections.get('Titulación', '')
    facultad = sections.get('Facultad/Escuela', '')
    periodo_docente = sections.get('Periodo docente', '')
    programa_formativo = load_programa_formativo()

    if not contenidos:
        return render(request, 'expand_content.html', {'error_message': 'No se encontró el contenido de la sección CONTENIDOS.'})


    

    # Definir el prompt para la API de IA
    prompt = (
        f"Aquí tienes el CONTENIDO de la asignatura {asignatura} de una titulación universitaria ({titulacion} que pertenece a la facultad de {facultad}). Esto te lo doy como contexto, no para que hagas referencia expresa.. Amplía el contenido actual, manteniendo la estructura, incluyendo subtemas o subapartados si es necesario. Céntrate sólo en el CONTENIDO, *NO incopores* otros componentes de la Guía Docente como la evaluación, bibliografía.... No incluyas nada antes ni después. "
        f"Contenido actual:\n{contenidos}\n\n"
        f"Es importante que este contenido vaya acorde al modelo pedagógico de la universidad:\n{programa_formativo}"
        f"Los títulos no los escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas)."
    )
    

    # Llamar a la API de GEMINI o ChatGPT para ampliar el contenido
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=1500,
            temperature=0.3,
        ),
    )

    # Contenido ampliado
    expanded_content = response.text

    # Renderizar el contenido original y ampliado
    return render(request, 'expand_content.html', {
        'original_content': contenidos,
        'expanded_content': expanded_content
    })


## EXPANDIR LAS ACTIVIDADES FORMATIVAS
def expand_activities_content(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el texto completo estructurado desde la sesión
    pdf_text = request.session.get('pdf_text', '')

    if not pdf_text:
        return render(request, 'expand_content.html', {'error_message': 'No se ha encontrado el texto del PDF.'})

    # Llamar a la función unificada para procesar y estructurar el contenido
    sections = process_and_parse_pdf_text(pdf_text)

    # Obtener las actividades formativas
    activities = sections.get('ACTIVIDADES FORMATIVAS', '')
    titulacion = sections.get('Titulación', '')
    facultad = sections.get('Facultad/Escuela', '')
    periodo_docente = sections.get('Periodo docente', '')
    programa_formativo = load_programa_formativo()

    if not activities:
        return render(request, 'expand_content.html', {'error_message': 'No se encontró la sección de ACTIVIDADES FORMATIVAS.'})

    # Definir el prompt para reformular y proponer nuevas actividades
    prompt = (
        f"Aquí tienes las actividades formativas de una Guía Docente universitaria. Amplíalas, manteniendo los apartados de docencia presencial y no presencial. "
        f"La titulación es {titulacion} de la Facultad {facultad} y los alumnos están en el {periodo_docente} semestre. Esto no hace falta que lo incluyas, es para que tengas contexto\n\n"
        f"Céntrate en las actividades formativas, no incluyas nada más de la Guía (como objetivos, contenidos, descripción...)\n\n"
        f"Actividades actuales:\n{activities}\n\n"
        f"Es importante que estas actividades vayan acordes al Programa Formativo de la universidad:\n{programa_formativo}."
        f"Los títulos no los escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas)."
    )

    # Llamar a la API para reformular y generar nuevas actividades
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=2000,
            temperature=0.4,
        ),
    )

    # Actividades reformuladas y nuevas
    expanded_content = response.text

    # Renderizar el contenido original y reformulado/propuesto
    return render(request, 'expand_activities_content.html', {
        'original_content': activities,
        'expanded_content': expanded_content
    })



## EXPANDIR LOS RESULTADOS DE APRENDIZAJE
def expand_resultados_aprendizaje_content(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el texto completo estructurado desde la sesión
    pdf_text = request.session.get('pdf_text', '')

    if not pdf_text:
        return render(request, 'expand_content.html', {'error_message': 'No se ha encontrado el texto del PDF.'})

    # Llamar a la función unificada para procesar y estructurar el contenido
    sections = process_and_parse_pdf_text(pdf_text)

    # Obtener las actividades formativas
    resultados_aprendizaje = sections.get('RESULTADOS DE APRENDIZAJE', '')
    competencias = sections.get('COMPETENCIAS', '')
    contenidos = sections.get('CONTENIDOS', '')
    titulacion = sections.get('Titulación', '')
    facultad = sections.get('Facultad/Escuela', '')
    periodo_docente = sections.get('Periodo docente', '')
    programa_formativo = load_programa_formativo()

    if not resultados_aprendizaje:
        return render(request, 'expand_content.html', {'error_message': 'No se encontró la sección de ACTIVIDADES FORMATIVAS.'})

    # Definir el prompt para reformular y proponer nuevas actividades
    prompt = (
        f"Aquí tienes los resultados de aprendizaje de una Guía Docente universitaria. Fíjándote en los {contenidos} y las {competencias}, amplía estos resultados de aprendizaje. Céntrate sólo en los RESULTADOS DE APRENDIZAJE, *NO incopores* otros componentes de la Guía Docente como la evaluación, bibliografía.... No incluyas nada antes ni después. "
        f"La titulación es {titulacion} de la Facultad {facultad} y los alumnos están en el {periodo_docente} semestre. Esto no hace falta que lo incluyas, es para que tengas contexto\n\n"
        f"Los resultados de aprendizaje actuales son:\n{resultados_aprendizaje}\n\n"
        f"Es importante que este resultados de aprendizaje vayan acordes al Programa Formativo de la universidad:\n{programa_formativo}"
        f"Los títulos no los escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas)."
    )

    # Llamar a la API de GEMINI o ChatGPT para reformular y generar nuevas actividades
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=1700,
            temperature=0.4,
        ),
    )

    # Actividades reformuladas y nuevas
    expanded_content = response.text

    # Renderizar el contenido original y reformulado/propuesto
    return render(request, 'expand_activities_content.html', {
        'original_content': resultados_aprendizaje,
        'expanded_content': expanded_content
    })

## EXPANDIR EL SISTEMA DE EVALUACIÓN DEL APRENDIZAJE
def expand_sistema_evaluacion_aprendizaje_content(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el texto completo estructurado desde la sesión
    pdf_text = request.session.get('pdf_text', '')

    if not pdf_text:
        return render(request, 'expand_content.html', {'error_message': 'No se ha encontrado el texto del PDF.'})

    # Llamar a la función unificada para procesar y estructurar el contenido
    sections = process_and_parse_pdf_text(pdf_text)

    # Obtener las actividades formativas
    sistema_evaluacion = sections.get('SISTEMA DE EVALUACIÓN DEL APRENDIZAJE', '')
    resultados_aprendizaje = sections.get('RESULTADOS DE APRENDIZAJE', '')
    asignatura = sections.get('Asignatura', '')
    competencias = sections.get('COMPETENCIAS', '')
    contenidos = sections.get('CONTENIDOS', '')
    titulacion = sections.get('Titulación', '')
    facultad = sections.get('Facultad/Escuela', '')
    periodo_docente = sections.get('Periodo docente', '')
    programa_formativo = load_programa_formativo()

    if not resultados_aprendizaje:
        return render(request, 'expand_content.html', {'error_message': 'No se encontró la sección de ACTIVIDADES FORMATIVAS.'})

    # Definir el prompt para reformular y proponer nuevas actividades
    prompt = (
        f"Aquí tienes el sistema de evaluación del aprendizaje de una Guía Docente universitaria. Fíjándote en los contenidos ({contenidos}(), las competencias ({competencias}) y los resultados de aprendizaje ({resultados_aprendizaje}) , amplía este sistema de evaluación dek aprendizaje. Céntrate sólo en el SISTEMA DE EVALUACIÓN DE APRENDIZAJE, *NO incopores* otros componentes de la Guía Docente como la evaluación, bibliografía.... No incluyas nada antes ni después. "
        f"La titulación es {titulacion} de la Facultad {facultad} y los alumnos están en el {periodo_docente} semestre. La asignatura se llama {asignatura} Esto no hace falta que lo incluyas, es para que tengas contexto\n\n"
        f"Sistema de evaluación del aprendizaje:\n{sistema_evaluacion}\n\n"
        f"Es importante que este sistema de evaluación del aprendizaje vaya acorde al Programa Formativo de la universidad:\n{programa_formativo}"
        f"Los títulos no los escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas)."
    )

    # Llamar a la API de GEMINI o ChatGPT para reformular y generar nuevas actividades
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=2000,
            temperature=0.4,
        ),
    )

    # Actividades reformuladas y nuevas
    expanded_content = response.text

    # Renderizar el contenido original y reformulado/propuesto
    return render(request, 'expand_sistemas_evaluacion_aprendizaje.html', {
        'original_content': sistema_evaluacion,
        'expanded_content': expanded_content
    })


## FORMAR PARA TRANSFORMAR - GENERAR ACTIVIDADES / PREGUNTAS / DEBATES

def formar_para_transformar(request):
    if not request.user.is_authenticated:
        return redirect('login')
    
    if request.method == 'POST':
        eleccion = request.POST.get('eleccion')
        numero = request.POST.get('numero')
        temas = request.POST.get('temas')
        version = request.POST.get('version')


        # Usar la versión seleccionada
        if version == 'modificada':
            contents = request.session.get('CONTENIDOS_MODIFICADOS', '')
        else:
            contents = request.session.get('CONTENIDOS', '')

        asignatura = request.session.get('Asignatura', '')
        print(f"Asignatura: {asignatura}")  
        periodo_docente = request.session.get('Periodo docente', '')
        titulacion = request.session.get('Titulación', '')
        facultad = request.session.get('Facultad/Escuela', '')
        competencias = request.session.get('COMPETENCIAS', '')
        resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')
        programa_formativo = load_programa_formativo()    # Obtener los temas de la sesión
        temas = request.session.get('themes', [])
        print(f"Temas en la sesión: {temas}")  # Verificar que los temas están en la sesión

        # Definir el prompt para la IA
        prompt = (
            f"Éste es el contenido ({contents}) de la Guía Docente de la asignatura ({asignatura}) en el ({periodo_docente}) "
            f"de la titulación ({titulacion}) que está en la Facultad ({facultad}). Los alumnos deben alcanzar estas competencias "
            f"({competencias}) y conseguir estos resultados de aprendizaje ({resultados_aprendizaje}). En la universidad seguimos "
            f"un modelo pedagógico resumido en {programa_formativo}. Todos estos datos te los doy como contexto, no hace falta que "
            f"los incorpores explícitamente en la respuesta. La tarea es proponer {numero} {eleccion} que haga relación a ese "
            f"modelo pedagógico, haciendo referencia expresa a los 3 verbos Despertar / Descubrir / Decidir. Estas {eleccion} deben estar relacionadas con el tema {temas}."
            f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas)."
        )

        # Incluir el tema seleccionado en el prompt si el usuario lo ha elegido
        if temas:
            prompt += f" Estas {eleccion} deben estar relacionadas con el tema {temas}."

        # Llamar a la API de GEMINI o ChatGPT para generar el contenido
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=3000,
                temperature=0.3,
            ),
        )

        # Contenido generado
        generated_content = response.text

        return render(request, 'formar_para_transformar_result.html', {
            'generated_content': generated_content
        })
    
    temas = request.session.get('temas', [])
    return render(request, 'formar_para_transformar.html', {'temas': temas})



## CALENDARIZACIÓN DE ACTIVIDADES
def calendarizacion(request):
    if not request.user.is_authenticated:
        return redirect('login')

    if request.method == 'POST':
        semanas = request.POST.get('semanas')
        días = request.POST.get('días')
        horas = request.POST.get('horas')

        try:
            semanas = int(semanas)  # Convertir semanas a entero
            días = int(días)      # Convertir días a entero     
            horas = int(horas)      # Convertir horas a entero
        except ValueError:
            return render(request, 'calendarizacion.html', {'error_message': 'Los valores de semanas y horas deben ser enteros.'})

        version = request.POST.get('version')

        # Usar la versión seleccionada
        if version == 'modificada':
            contents = request.session.get('CONTENIDOS_MODIFICADOS', '')
        else:
            contents = request.session.get('CONTENIDOS', '')

        asignatura = request.session.get('Asignatura', '')
        periodo_docente = request.session.get('Periodo docente', '')
        titulacion = request.session.get('Titulación', '')
        facultad = request.session.get('Facultad/Escuela', '')
        competencias = request.session.get('COMPETENCIAS', '')
        resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')
        programa_formativo = load_programa_formativo()    # Obtener los temas de la sesión
        contenidos = request.session.get('CONTENIDOS', '')
        temas = request.session.get('temas', [])
        # print(f"Temas en la sesión: {temas}")  # Verificar que los temas están en la sesión

        # Definir el prompt para la IA
        prompt = (
            f"Éste es el contenido ({contenidos}) de la Guía Docente de la asignatura ({asignatura}) en el ({periodo_docente}) "
            f"de la titulación ({titulacion}) que está en la Facultad ({facultad}). Todos estos datos te los doy como contexto, no hace falta que "
            f"los incorpores explícitamente en la respuesta. La tarea es proponer al profesor una posible calendarización de los contenidos "
            f"en las {semanas} semanas que dura el curso, sabiendo imparte clases {días} cada semana, para un total de {horas} horas / semanales. Incluye algún día para un examen parcial "
            f"y al menos uno al final para exposición de trabajos finales. Da el resultado en formato tabla con los siguientes encabezados: SEMANA | Nº CLASE DE ESA SEMANA | TEMA | CONTENIDO | TIPO DE ACTIVIDAD | TIEMPO PREVISTO"
        )

        # Llamar a la API de GEMINI o ChatGPT para generar la calendarización
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=1500,
                temperature=0.3,
            ),
        )

        # Contenido generado
        generated_calendar = response.text
        # print(f"Calendarización generada: {generated_calendar}")

        # Renderizar el resultado en una nueva página
        return render(request, 'calendarizacion_result.html', {
            'generated_calendar': generated_calendar
        })

    return render(request, 'calendarizacion.html')


##RÚBRICA DE EVALUACIÓN

def generar_rubrica_evaluacion(request):
    if request.method == 'POST':
        # Obtener los datos del formulario
        tipo_actividad = request.POST.get('tipo_actividad')
        if tipo_actividad == 'Otro':
            tipo_actividad = request.POST.get('otro_actividad')

        tema = request.POST.get('tema_descripcion')
        n_indicadores = int(request.POST.get('n_indicadores'))
        n_niveles = int(request.POST.get('n_niveles'))

        # Información de contexto desde la Guía Docente
        contenido = "Aquí va el contenido de la Guía Docente"
        asignatura = "Asignatura"
        periodo_docente = "Periodo docente"
        titulacion = "Titulación"
        facultad = "Facultad"
        competencias = "Competencias"
        resultados_aprendizaje = "Resultados de aprendizaje"
        programa_formativo = "Programa formativo"

        # Definir el prompt para la API
        prompt = (
            f"Éste es el contenido ({contenido}) de la Guía Docente de la asignatura ({asignatura}) "
            f"en el ({periodo_docente}) de la titulación ({titulacion}) que está en la Facultad ({facultad}). "
            f"Los alumnos deben alcanzar estas competencias ({competencias}) y conseguir estos resultados de aprendizaje ({resultados_aprendizaje}). "
            f"En la universidad seguimos un modelo pedagógico resumido en {programa_formativo}. "
            f"Todos estos datos te los doy como contexto, no hace falta que los incorpores explícitamente en la respuesta. "
            f"La tarea es proponer una rúbrica en formato tabla con {n_indicadores} indicadores y cada uno de ellos debe tener {n_niveles} niveles, "
            f"que evalúe {tipo_actividad} sobre {tema}."
        )

        # Llamada a la API de GEMINI o GPT
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=2000,
                temperature=0.4,
            ),
        )

        # Log de la respuesta
        print("Respuesta de la API:", response.text)

        # Procesar la respuesta para construir una lista de diccionarios
        rubrica_generada = response.text.split("\n")
        rubrica = []

        # Filtrar solo las líneas relevantes para la rúbrica
        for line in rubrica_generada:
            if '|' in line:  # Filtrar solo las líneas con estructura de tabla
                parts = line.split("|")

                # Remover columnas vacías
                parts = [part.strip() for part in parts if part.strip()]

                indicador = parts[0]  # El primer elemento es el indicador
                niveles = parts[1:]  # El resto son los niveles

                # Asegurar que los niveles coincidan con la cantidad esperada
                if len(niveles) < n_niveles:
                    niveles.extend(['---'] * (n_niveles - len(niveles)))
                elif len(niveles) > n_niveles:
                    niveles = niveles[:n_niveles]  # Truncar si hay más niveles de los esperados

                # Agregar la fila procesada
                rubrica.append({'indicador': indicador, 'niveles': niveles})

        # Guardar la rúbrica en la sesión
        request.session['rubrica'] = rubrica

        # Renderizar el resultado en una nueva página
        return render(request, 'rubrica_resultados.html', {
            'rubrica': rubrica,
            'range_niveles': range(1, n_niveles + 1)
        })

    return render(request, 'rubrica_intermedia.html')




def rubrica_pagina_intermedia(request):
    if not request.user.is_authenticated:
        return redirect('login')

    
    return render(request, 'rubrica_intermedia.html')



#### APLICACIONES CON APUNTES

import docx  # Para manejar documentos de Word
import PyPDF2
from django.core.files.storage import FileSystemStorage

# Función para extraer apuntes de PDF
def extract_apuntes_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        apuntes = ""
        for page_num in range(len(reader.pages)):
            apuntes += reader.pages[page_num].extract_text()
    return apuntes

# Función para extraer apuntes de Word
def extract_apuntes_from_word(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

# Vista para subir y procesar apuntes
def upload_apuntes_view(request):
    if request.method == 'POST' and request.FILES.getlist('apuntes'):
        apuntes = request.FILES.getlist('apuntes')
        apuntes_texts = []

        for apunte in apuntes:
            fs = FileSystemStorage(location=settings.MEDIA_ROOT)
            filename = fs.save(apunte.name, apunte)

            # Determinar el tipo de archivo (PDF o Word)
            if apunte.name.endswith('.pdf'):
                pdf_path = os.path.join(settings.MEDIA_ROOT, filename)
                text = extract_apuntes_from_pdf(pdf_path)
                apuntes_texts.append({'name': apunte.name, 'content': text})
            elif apunte.name.endswith('.doc') or apunte.name.endswith('.docx'):
                word_path = os.path.join(settings.MEDIA_ROOT, filename)
                text = extract_apuntes_from_word(word_path)
                apuntes_texts.append({'name': apunte.name, 'content': text})

        # Guardar el contenido procesado en la sesión para usarlo más adelante
        request.session['apuntes_texts'] = apuntes_texts

        # Redirigir al "dashboard_apuntes"
        return redirect('dashboard_apuntes')

    return render(request, 'upload_apuntes.html')

# Vista del dashboard de apuntes
def dashboard_apuntes_view(request):
    apuntes_texts = request.session.get('apuntes_texts', [])
    return render(request, 'dashboard_apuntes.html', {'apuntes_texts': apuntes_texts})

## APLICACIÓN PUNTOS CLAVE DE APUNTES

def puntos_clave_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener los apuntes desde la sesión
    apuntes_text = request.session.get('apuntes_texts', [])

    if not apuntes_text:
        return render(request, 'puntos_clave.html', {'error_message': 'No se ha encontrado el texto de los apuntes.'})

    # Obtener datos de contexto (asignatura, titulación, facultad, programa formativo)
    asignatura = request.session.get('Asignatura', '')
    titulacion = request.session.get('Titulación', '')
    facultad = request.session.get('Facultad/Escuela', '')
    programa_formativo = load_programa_formativo()

    # Comprobar si existe el contenido de los apuntes
    if not apuntes_text:
        return render(request, 'puntos_clave.html', {'error_message': 'No se encontró el contenido de los apuntes.'})

    # Definir el prompt para extraer los puntos clave
    prompt = (
        f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}. "
        f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
        f"Todo esto no hace falta que lo incluyas, es para que tengas contexto.\n"
        f"Quiero que del contenido de los apuntes extraigas el contenido clave, lo más didáctico posible. Una especie de resumen ejecutivo, pero enfocado a la docencia. . Aunque no haga falta que sea explícito, ten en cuenta el modelo pedagógico"
        f"Será un documento para el profesor.\n"
        f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
        f"El contenido de los apuntes es:\n{apuntes_text}"
    )

    # Llamar a la API de GEMINI o ChatGPT para generar los puntos clave
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=3000,
            temperature=0.4,
        ),
    )

    # Puntos clave generados por la IA
    puntos_clave = response.text

    # Renderizar el contenido original y los puntos clave generados
    return render(request, 'puntos_clave.html', {
        'original_content': apuntes_text,
        'puntos_clave': puntos_clave
    })


## APLICACIÓN DE FAQs CON LOS APUNTES

def faqs_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener los apuntes desde la sesión
    apuntes_text = request.session.get('apuntes_texts', [])

    if not apuntes_text:
        return render(request, 'faqs.html', {'error_message': 'No se ha encontrado el texto de los apuntes.'})

    # Obtener datos de contexto (asignatura, titulación, facultad, programa formativo)
    asignatura = request.session.get('Asignatura', '')
    titulacion = request.session.get('Titulación', '')
    facultad = request.session.get('Facultad/Escuela', '')
    programa_formativo = load_programa_formativo()

    # Definir el prompt para extraer las FAQs
    prompt = (
        f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}. "
        f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
        f"Todo esto no hace falta que lo incluyas, es para que tengas contexto.\n"
        f"Quiero que del contenido de los apuntes extraigas preguntas a desarrollar para poner en un examen, un mínimo de 5 y un máximo de 10. "
        f"Sólo escribe las preguntas, nada delante ni detrás"
        f"Preguntas clave que tengan que ver con los apuntes y el modelo pedagógico.\n"
        f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
        f"El contenido de los apuntes es:\n{apuntes_text}"
    )

    # Llamar a la API de GEMINI o ChatGPT para generar las FAQs
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=3000,
            temperature=0.4,
        ),
    )

    # FAQs generadas por la IA
    faqs_content = response.text

    # Renderizar el contenido original y las FAQs generadas
    return render(request, 'faqs_apuntes.html', {
        'original_content': apuntes_text,
        'faqs_content': faqs_content
    })


## APLICACIÓN GLOSARIO DE TÉRMINOS CON LOS APUNTES

def glosario_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener los apuntes desde la sesión
    apuntes_text = request.session.get('apuntes_texts', '')

    if not apuntes_text:
        return render(request, 'glosario.html', {'error_message': 'No se ha encontrado el texto de los apuntes.'})

    # Obtener datos de contexto (asignatura, titulación, facultad, programa formativo)
    asignatura = request.session.get('Asignatura', '')
    titulacion = request.session.get('Titulación', '')
    facultad = request.session.get('Facultad/Escuela', '')
    programa_formativo = load_programa_formativo()

    # Definir el prompt para extraer el glosario de términos
    prompt = (
        f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}. "
        f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
        f"Todo esto no hace falta que lo incluyas, es para que tengas contexto.\n"
        f"Quiero que del contenido de los apuntes extraigas un glosario de entre 15 y 25 términos clave. Aunque no haga falta que sea explícito, ten en cuenta el modelo pedagógico\n"
        f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
        f"El contenido de los apuntes es:\n{apuntes_text}"
    )

    # Llamar a la API de GEMINI para generar el glosario
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=3000,
            temperature=0.4,
        ),
    )

    # Glosario generado por la IA
    glosario_content = response.text

    # Renderizar el contenido original y el glosario generado
    return render(request, 'glosario_apuntes.html', {
        'original_content': apuntes_text,
        'glosario_content': glosario_content
    })


## CASO PRÁCTICO CON LOS APUNTES

def caso_practico_view(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener los apuntes desde la sesión
    apuntes_text = request.session.get('apuntes_texts', '')

    if not apuntes_text:
        return render(request, 'caso_practico.html', {'error_message': 'No se ha encontrado el texto de los apuntes.'})

    # Obtener datos de contexto (asignatura, titulación, facultad, programa formativo)
    asignatura = request.session.get('Asignatura', '')
    titulacion = request.session.get('Titulación', '')
    facultad = request.session.get('Facultad/Escuela', '')
    programa_formativo = load_programa_formativo()
    contenidos = request.session.get('CONTENIDOS', '')
    resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')

    # Definir el prompt para generar el caso práctico
    prompt = (
        f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}. "
        f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
        f"Todo esto no hace falta que lo incluyas, es para que tengas contexto.\n"
        f"Quiero que del contenido de los apuntes generes un caso práctico, completo, inventado, incluyendo el contexto, personajes, que sea atractivo, relevante, desafiante y resolubre. "
        f"A modo de guía para el profesor, indica el objeto del caso, si se necesitan lecturas adicionales y cualquier otra información que consideres.\n"
        f"El esquema sería:\n"
        f"PLANTEAMIENTO DEL CASO\nMATERIAL\nOBSERVACIONES PARA EL PROFESOR\n"
        f"Aunque no haga falta que sea explícito, ten en cuenta el modelo pedagógico.\n"
        f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
        f"El contenido de la asignatura es {contenidos}\n"
        f"Y sus resultados de aprendizaje son: {resultados_aprendizaje}\n"
        f"El contenido de los apuntes es:\n{apuntes_text}"
    )

    # Llamar a la API de GEMINI para generar el caso práctico
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            max_output_tokens=5000,
            temperature=0.4,
        ),
    )

    # Caso práctico generado por la IA
    caso_practico_content = response.text

    # Renderizar el contenido original y el caso práctico generado
    return render(request, 'caso_practico_apuntes.html', {
        'original_content': apuntes_text,
        'caso_practico_content': caso_practico_content
    })


## APLICACIÓN DE PREGUNTAS TIPO TEST CON LOS APUNTES

import logging

# Configurar logger básico
logging.basicConfig(level=logging.DEBUG)

def tipo_test_view(request):
    if not request.user.is_authenticated:
        logging.debug("Usuario no autenticado. Redirigiendo a login.")
        return redirect('login')

    if request.method == 'POST':
        logging.debug("Petición POST recibida en tipo_test_view")

        # Obtener los parámetros del formulario
        numero = request.POST.get('numero')
        opciones = request.POST.get('opciones')
        respuesta_unica = request.POST.get('respuesta_unica')

        logging.debug(f"Datos recibidos: número = {numero}, opciones = {opciones}, respuesta_unica = {respuesta_unica}")

        # Asegurarse de que los valores no sean nulos o vacíos
        if not numero or not opciones or not respuesta_unica:
            logging.debug("Faltan datos en el formulario")
            return render(request, 'tipo_test_apuntes.html', {
                'error_message': 'Por favor, rellena todos los campos.'
            })

        tipo_respuesta = 'sólo haya una correcta' if respuesta_unica == 'unica' else 'pueda haber más de una correcta'

        # Validación básica de los inputs del formulario
        try:
            numero = int(numero)
            opciones = int(opciones)
            if numero < 1 or numero > 10 or opciones < 2 or opciones > 7:
                raise ValueError("Valores fuera del rango permitido")
            logging.debug("Validación de número y opciones correcta.")
        except ValueError:
            logging.debug("Error en la validación de los inputs del formulario")
            return render(request, 'tipo_test_apuntes.html', {
                'error_message': 'Por favor, selecciona valores válidos para el número de preguntas y opciones.'
            })

        # Obtener los apuntes y el contexto desde la sesión
        apuntes_text = request.session.get('apuntes_texts', '')
        asignatura = request.session.get('Asignatura', '')
        titulacion = request.session.get('Titulación', '')
        facultad = request.session.get('Facultad/Escuela', '')
        programa_formativo = load_programa_formativo()
        contenidos = request.session.get('CONTENIDOS', '')
        resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')

        logging.debug(f"Apuntes cargados: {bool(apuntes_text)}, Asignatura: {asignatura}, Titulación: {titulacion}, Facultad: {facultad}")

        # Comprobar que los apuntes estén en la sesión
        if not apuntes_text:
            logging.debug("Apuntes no encontrados en la sesión.")
            return render(request, 'tipo_test_apuntes.html', {
                'error_message': 'No se han encontrado apuntes. Por favor, sube los apuntes antes de generar las preguntas.'
            })

        # Definir el prompt para la generación de preguntas tipo test
        prompt = (
            f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}. "
            f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
            f"Todo esto no hace falta que lo incluyas, es para que tengas contexto.\n"
            f"Quiero que del contenido de los apuntes generes {numero} preguntas tipo test, de {opciones} opciones, en las que {tipo_respuesta}.\n"
            f"La respuesta correcta debe ir marcada con un asterisco delante.\n"
            f"Algunos tips a seguir para hacer buenas preguntas tipo test:\n. Cada ítem debe reflejar un contenido específico\n. Utiliza una redacción novedosa para las preguntas. Evita copias literales o paráfrasis muy semejantes\n. Evita preguntas basadas en la opinión\n. Evita preguntas trampa\n. Utiliza un vocabulario simple\n. Minimiza el tiempo de lectura\n. Incluye la idea principal en el enunciado y no en las opciones\n. Redacta la pregunta en positivo y, si debes usar negativas como «no» o «excepto», escríbelas en mayúscula\n. Las opciones de respuesta no deben solaparse\n. Redacta homogéneamente las opciones de respuesta\n. Usa errores típicos para crear distractores, aprovechando los errores habituales que cometen los alumnos sobre un tema para introducirlos como opciones incorrectas."
            f"Sólo incluye las preguntas y respuestas, no pongas nada ni delante ni detrás.\n"
            f"NO formatees en markdown, pon las preguntas en texto plano"
            f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
            f"El contenido de los apuntes es:\n{apuntes_text}\n"
            f"El contenido de la asignatura es {contenidos}\n"
            f"Y sus resultados de aprendizaje son: {resultados_aprendizaje}"
        )

        logging.debug("Prompt generado correctamente. Haciendo la llamada a la API...")

        # Llamada a la API de GEMINI para generar las preguntas tipo test
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=2000,
                temperature=0.4,
            ),
        )

        # Actividades reformuladas y nuevas
        preguntas_respuestas = response.text
        logging.debug(f"Respuesta de la API recibida: {preguntas_respuestas[:500]}")  # Muestra solo los primeros 500 caracteres

        # Guardar en la sesión para exportación o futuras vistas
        request.session['preguntas_respuestas'] = preguntas_respuestas

        # Renderizar las preguntas generadas
        return render(request, 'tipo_test_resultado.html', {
            'preguntas_respuestas': preguntas_respuestas,
        })

    # Si no es POST, renderiza la página inicial
    logging.debug("Renderizando tipo_test_apuntes.html (GET request)")
    return render(request, 'tipo_test_apuntes.html')


## ACTIVIDADES DE EVALUACIÓN CON LOS APUNTES

def extraer_temas_apuntes(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener el contenido de los apuntes desde la sesión (apuntes_texts es una lista de diccionarios)
    apuntes_texts = request.session.get('apuntes_texts', [])
    
    if not apuntes_texts:
        print("No se han encontrado los apuntes en la sesión.")
        return render(request, 'pagina_intermedia_actividades.html', {
            'error_message': 'No se han encontrado los apuntes.'
        })

    # Concatenar el contenido de todos los apuntes en un solo texto
    apuntes_completos = ' '.join([apunte['content'] for apunte in apuntes_texts])
    print(f"Texto de apuntes concatenado: {apuntes_completos[:500]}")  # Mostrar los primeros 500 caracteres

    # Llamada a la API para extraer los temas principales
    prompt = (
        f"Estos son los apuntes de una asignatura universitaria; "
        f"extrae un listado con los temas principales que trata, separados por '|'.\n\n"
        f"{apuntes_completos}"
    )
    print(f"Prompt enviado a la API: {prompt[:500]}")  # Mostrar los primeros 500 caracteres del prompt

    try:
        # Llamada a la API de GEMINI o la que estés usando
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=500,
                temperature=0.0,
            ),
        )

        # Procesar la respuesta y extraer los temas
        temas_apuntes = response.text.split('|')
        temas_apuntes = [tema.strip() for tema in temas_apuntes if tema.strip()]

        print(f"Temas extraídos: {temas_apuntes}")

        if not temas_apuntes:
            print("No se han encontrado temas en los apuntes.")
            return render(request, 'pagina_intermedia_actividades.html', {
                'error_message': 'No se han encontrado temas en los apuntes.'
            })

        # Guardar los temas en la sesión
        request.session['temas_apuntes'] = temas_apuntes
        print("Temas guardados en la sesión.")

    except Exception as e:
        print(f"Error en la llamada a la API: {e}")
        return render(request, 'pagina_intermedia_actividades.html', {
            'error_message': f'Error al procesar los apuntes: {str(e)}'
        })

    # Redirigir a la página intermedia
    return redirect('pagina_intermedia_actividades')




def actividades_pagina_intermedia_apuntes(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Comprobar si se obtienen los temas de la sesión
    temas_apuntes = request.session.get('temas_apuntes', [])
    print(f"Temas disponibles en la sesión: {temas_apuntes}")

    if not temas_apuntes:
        print("No se han encontrado temas en la sesión.")
        return render(request, 'pagina_intermedia_actividades.html', {
            'error_message': 'No se han encontrado temas en los apuntes.'
        })

    return render(request, 'pagina_intermedia_actividades.html', {
        'temas': temas_apuntes
    })



def actividades_pagina_intermedia(request):
    if not request.user.is_authenticated:
        return redirect('login')

    # Obtener los temas extraídos de la sesión
    temas_apuntes = request.session.get('temas_apuntes', [])

    if not temas_apuntes:
        return render(request, 'pagina_intermedia_actividades.html', {
            'error_message': 'No se han encontrado los temas extraídos de los apuntes.'
        })

    # Renderizar la página intermedia donde se seleccionan las opciones
    return render(request, 'pagina_intermedia_actividades.html', {
        'temas_apuntes': temas_apuntes
    })

def generar_actividades_evaluacion(request):
    if not request.user.is_authenticated:
        return redirect('login')

    if request.method == 'POST':
        numero_actividades = request.POST.get('numero_actividades')
        tema_apuntes = request.POST.get('tema_apuntes')

        # Obtener contexto adicional de la sesión
        apuntes_text = request.session.get('apuntes_texts', '')
        asignatura = request.session.get('Asignatura', '')
        titulacion = request.session.get('Titulación', '')
        facultad = request.session.get('Facultad/Escuela', '')
        programa_formativo = load_programa_formativo()
        contenidos = request.session.get('CONTENIDOS', '')
        competencias = request.session.get('COMPETENCIAS', '')
        resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')

        # Definir el prompt para generar actividades de evaluación
        prompt = (
            f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}.\n"
            f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
            f"El contenido de la asignatura es {contenidos}\n"
            f"Las competencias que tiene que adquirir son: {competencias}\n"
            f"Y sus resultados de aprendizaje son: {resultados_aprendizaje}\n"
            f"Todo esto no hace falta que lo incluyas explícitamente, es para que tengas contexto.\n"
            f"Quiero que del contenido de los apuntes generes 1 actividad de evaluación sobre el tema '{tema_apuntes}' "
            f"para reflejar que el alumno ha conseguido las competencias y resultados de aprendizaje. Indica la actividad, desarrollo, y cómo se demuestra que el alumno ha alcanzado los resultados de aprendizaje (cuáles se evalúan y cómo).\n"
            f"Es necesario que aparezca cómo se relaciona esta actividad con el modelo pedagógico, específicamente con los 3 verbos\n"
            f"Está PROHIBIDO que la actividad propuesta sea el análisis de un caso: tiene que ser algo diferente\n\n"
            f"NO escribas en versalita: usa siempre el formato oración (primera palabra en mayúsculas, resto en minúsculas).\n"
            f"El contenido de los apuntes es:\n{apuntes_text}"
            f"Está PROHIBIDO que la actividad propuesta sea el análisis de un caso: tiene que ser algo diferente\n\n"
        )

        # Llamada a la API de GEMINI para generar las actividades de evaluación
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=6000,
                temperature=1,
            ),
        )

        actividades_evaluacion = response.text

        return render(request, 'actividades_evaluacion_resultado.html', {
            'actividades_evaluacion': actividades_evaluacion,
        })



def generar_actividades_evaluacion(request):
    if not request.user.is_authenticated:
        return redirect('login')

    if request.method == 'POST':
        numero_actividades = request.POST.get('numero_actividades')
        tema_apuntes = request.POST.get('tema')

        # Obtener contexto adicional de la sesión
        apuntes_text = request.session.get('apuntes_texts', '')
        asignatura = request.session.get('Asignatura', '')
        titulacion = request.session.get('Titulación', '')
        facultad = request.session.get('Facultad/Escuela', '')
        programa_formativo = load_programa_formativo()
        contenidos = request.session.get('CONTENIDOS', '')
        competencias = request.session.get('COMPETENCIAS', '')
        resultados_aprendizaje = request.session.get('RESULTADOS DE APRENDIZAJE', '')

        # Definir el prompt para generar actividades de evaluación
        prompt = (
            f"Vas a recibir los apuntes de la asignatura {asignatura}, que pertenece a la titulación {titulacion} de la Facultad {facultad}.\n"
            f"La universidad sigue un modelo pedagógico, que es el siguiente:\n{programa_formativo}\n"
            f"El contenido de la asignatura es {contenidos}\n"
            f"Las competencias que tiene que adquirir son: {competencias}\n"
            f"Y sus resultados de aprendizaje son: {resultados_aprendizaje}\n"
            f"Todo esto no hace falta que lo incluyas explícitamente, es para que tengas contexto.\n"
            f"Quiero que del contenido de los apuntes generes {numero_actividades} actividad(es) de evaluación sobre el tema '{tema_apuntes}' "
            f"para reflejar que el alumno ha conseguido las competencias y resultados de aprendizaje. Indica la actividad, desarrollo, y cómo se demuestra que el alumno ha alcanzado los resultados de aprendizaje (cuáles se evalúan y cómo).\n\n"
            f"El contenido de los apuntes es:\n{apuntes_text}"
        )

        # Llamada a la API de GEMINI para generar las actividades de evaluación
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=2000,
                temperature=0.4,
            ),
        )

        actividades_evaluacion = response.text

        return render(request, 'actividades_evaluacion_resultado.html', {
            'actividades_evaluacion': actividades_evaluacion,
        })


### AQUÍ ACABAN LAS APLICACIONES


def load_programa_formativo():
    with open(os.path.join(settings.BASE_DIR, 'programa_formativo.txt'), 'r', encoding='utf-8') as f:
        return f.read()


# Añadimos una función para generar y descargar un documento Word con la descripción reformulada
from django.http import HttpResponse
from docx import Document

def download_reformulated_description(request):
    if not request.user.is_authenticated:
        return redirect('login')

    reformulated_text = request.session.get('reformulated_text', '')

    document = Document()
    document.add_heading('Descripción Reformulada de la Asignatura', level=1)
    document.add_paragraph(reformulated_text)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=descripcion_reformulada.docx'
    document.save(response)
    return response

from xml.etree.ElementTree import Element, SubElement, tostring
from django.http import HttpResponse

def descargar_qti_view(request):
    preguntas_respuestas = request.session.get('preguntas_respuestas', '')

    if not preguntas_respuestas:
        return HttpResponse("No se encontraron preguntas generadas.", content_type="text/plain")

    # Crear el árbol XML en formato QTI 2.1
    assessment = Element('assessmentTest', attrib={'xmlns': 'http://www.imsglobal.org/xsd/imsqti_v2p1'})

    # Dividir las preguntas generadas por líneas
    preguntas = preguntas_respuestas.split('\n\n')

    for i, pregunta in enumerate(preguntas, start=1):
        item = SubElement(assessment, 'assessmentItem', attrib={'identifier': f'q{i}', 'title': f'Pregunta {i}'})
        body = SubElement(item, 'itemBody')
        question_text = pregunta.split('\n')[0]  # Primera línea: el enunciado
        question = SubElement(body, 'p')
        question.text = question_text

        choices = pregunta.split('\n')[1:]  # Opciones de respuesta
        for choice in choices:
            answer = SubElement(body, 'choiceInteraction')
            answer.text = choice

    # Convertir el árbol XML a bytes
    xml_data = tostring(assessment)

    # Crear la respuesta HTTP con el archivo XML adjunto
    response = HttpResponse(xml_data, content_type='application/xml')
    response['Content-Disposition'] = 'attachment; filename="preguntas.qti.xml"'
    return response


import openpyxl
from django.http import HttpResponse



def clean_markdown_rubrica(text):
    # Remover negritas y títulos de Markdown (**, ##, etc.)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Elimina negritas
    text = re.sub(r'##\s*(.*?)\s*', r'\1', text)  # Elimina títulos con ##
    text = re.sub(r'\*\s*(.*?)\s*', r'\1', text)  # Elimina listas con *
    return text


def descargar_rubrica_xlsx(request):
    import openpyxl
    from openpyxl.utils import get_column_letter
    from django.http import HttpResponse

    # Obtener la rúbrica de la sesión
    rubrica = request.session.get('rubrica')

    if not rubrica:
        return HttpResponse("No hay rúbrica para descargar.")

    # Crear un archivo Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Rúbrica"

    # Escribir encabezados
    headers = ["Indicador"] + [f'Nivel {i+1}' for i in range(len(rubrica[0]['niveles']))]
    ws.append(headers)

    # Escribir cada fila de la rúbrica, aplicando clean_markdown
    for row in rubrica:
        cleaned_indicator = clean_markdown_rubrica(row['indicador'])
        cleaned_niveles = [clean_markdown_rubrica(nivel) for nivel in row['niveles']]
        ws.append([cleaned_indicator] + cleaned_niveles)

    # Ajustar el ancho de las columnas automáticamente
    for col in ws.columns:
        max_length = max(len(str(cell.value)) for cell in col)
        column_letter = get_column_letter(col[0].column) 
        ws.column_dimensions[column_letter].width = max_length

    # Crear una respuesta HTTP con el archivo Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=rubrica.xlsx'
    wb.save(response)

    return response
